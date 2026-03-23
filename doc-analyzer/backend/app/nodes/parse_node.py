"""
文档解析节点
支持 PDF、DOCX、TXT、URL
"""
import os
import re
import requests
from typing import Dict, Any
from bs4 import BeautifulSoup


def parse_document(file_path: str, file_type: str) -> Dict[str, Any]:
    """
    解析文档为纯文本
    
    Args:
        file_path: 文件路径或 URL
        file_type: 文件类型 (pdf, docx, txt, url)
    
    Returns:
        {
            "text": "解析后的纯文本",
            "title": "文档标题（如有）",
            "page_count": 页数（PDF）,
            "word_count": 字数统计
        }
    """
    if file_type == 'pdf':
        return _parse_pdf(file_path)
    elif file_type == 'docx':
        return _parse_docx(file_path)
    elif file_type == 'txt':
        return _parse_txt(file_path)
    elif file_type == 'url':
        return _parse_url(file_path)
    else:
        raise ValueError(f"不支持的文件类型: {file_type}")


def _parse_pdf(file_path: str) -> Dict[str, Any]:
    """解析 PDF 文件"""
    try:
        import pdfplumber
    except ImportError:
        raise ImportError("请安装 pdfplumber: pip install pdfplumber")
    
    text_parts = []
    title = ""
    page_count = 0
    
    with pdfplumber.open(file_path) as pdf:
        page_count = len(pdf.pages)
        
        # 尝试从第一页提取标题
        if page_count > 0:
            first_page_text = pdf.pages[0].extract_text()
            if first_page_text:
                lines = first_page_text.strip().split('\n')
                if lines:
                    title = lines[0].strip()[:100]  # 取第一行作为标题
        
        # 提取所有页面文本
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)
    
    full_text = '\n'.join(text_parts)
    
    return {
        "text": full_text,
        "title": title,
        "page_count": page_count,
        "word_count": len(full_text)
    }


def _parse_docx(file_path: str) -> Dict[str, Any]:
    """解析 Word 文档"""
    try:
        from docx import Document
    except ImportError:
        raise ImportError("请安装 python-docx: pip install python-docx")
    
    doc = Document(file_path)
    
    # 提取段落
    text_parts = []
    for para in doc.paragraphs:
        if para.text.strip():
            text_parts.append(para.text)
    
    # 提取表格文本
    for table in doc.tables:
        for row in table.rows:
            row_text = ' '.join([cell.text for cell in row.cells if cell.text.strip()])
            if row_text:
                text_parts.append(row_text)
    
    full_text = '\n'.join(text_parts)
    
    # 尝试获取标题（第一个非空段落）
    title = ""
    for para in doc.paragraphs:
        if para.text.strip():
            title = para.text.strip()[:100]
            break
    
    return {
        "text": full_text,
        "title": title,
        "page_count": 0,
        "word_count": len(full_text)
    }


def _parse_txt(file_path: str) -> Dict[str, Any]:
    """解析文本文件"""
    # 尝试多种编码
    encodings = ['utf-8', 'gbk', 'gb2312', 'utf-16', 'latin-1']
    
    for encoding in encodings:
        try:
            with open(file_path, 'r', encoding=encoding) as f:
                text = f.read()
            break
        except UnicodeDecodeError:
            continue
    else:
        raise ValueError(f"无法解码文件: {file_path}")
    
    # 获取标题（第一行）
    lines = text.strip().split('\n')
    title = lines[0][:100] if lines else ""
    
    return {
        "text": text,
        "title": title,
        "page_count": 0,
        "word_count": len(text)
    }


def _parse_url(url: str) -> Dict[str, Any]:
    """解析网页 URL"""
    try:
        headers = {
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36'
        }
        response = requests.get(url, headers=headers, timeout=30)
        response.raise_for_status()
        
        # 自动检测编码
        response.encoding = response.apparent_encoding
        
        soup = BeautifulSoup(response.text, 'html.parser')
        
        # 移除脚本和样式
        for script in soup(["script", "style", "nav", "footer", "header"]):
            script.decompose()
        
        # 获取标题
        title = ""
        if soup.title:
            title = soup.title.string.strip() if soup.title.string else ""
        
        # 获取正文
        # 优先查找 article 或 main 标签
        article = soup.find('article') or soup.find('main')
        if article:
            text = article.get_text(separator='\n', strip=True)
        else:
            # 查找内容最多的 div
            paragraphs = soup.find_all('p')
            text_parts = [p.get_text(strip=True) for p in paragraphs if len(p.get_text(strip=True)) > 20]
            text = '\n'.join(text_parts)
        
        # 清理多余空白
        text = re.sub(r'\n+', '\n', text)
        text = re.sub(r' +', ' ', text)
        
        return {
            "text": text,
            "title": title[:100],
            "page_count": 0,
            "word_count": len(text)
        }
        
    except requests.RequestException as e:
        raise ValueError(f"请求失败: {str(e)}")
